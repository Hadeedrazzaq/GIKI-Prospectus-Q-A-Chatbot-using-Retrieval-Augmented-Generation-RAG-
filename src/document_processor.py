"""
Document processing module for GIKI Prospectus Q&A Chatbot
Handles PDF, DOCX, and TXT file extraction and text processing
"""
import os
import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Optional, Tuple
import logging
from pathlib import Path

from config import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    def __init__(self):
        self.allowed_extensions = Config.ALLOWED_EXTENSIONS
        self.max_file_size = Config.MAX_FILE_SIZE
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate uploaded file"""
        try:
            file_path = Path(file_path)
            
            # Check file extension
            if file_path.suffix.lower() not in self.allowed_extensions:
                return False, f"File type {file_path.suffix} not supported. Allowed: {', '.join(self.allowed_extensions)}"
            
            # Check file size
            if file_path.stat().st_size > self.max_file_size:
                return False, f"File size exceeds maximum limit of {self.max_file_size // (1024*1024)}MB"
            
            # Check if file exists
            if not file_path.exists():
                return False, "File does not exist"
            
            return True, "File is valid"
            
        except Exception as e:
            return False, f"Error validating file: {str(e)}"
    
    def extract_text_from_pdf(self, file_path: str) -> List[Dict[str, any]]:
        """Extract text from PDF file with page information"""
        try:
            doc = fitz.open(file_path)
            pages = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():  # Only add non-empty pages
                    pages.append({
                        'page_number': page_num + 1,
                        'text': text.strip(),
                        'file_name': os.path.basename(file_path)
                    })
            
            doc.close()
            logger.info(f"Extracted {len(pages)} pages from PDF: {file_path}")
            return pages
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> List[Dict[str, any]]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            # Combine all text
            full_text = '\n'.join(text_content)
            
            if full_text.strip():
                return [{
                    'page_number': 1,  # DOCX doesn't have explicit page numbers
                    'text': full_text,
                    'file_name': os.path.basename(file_path)
                }]
            else:
                return []
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> List[Dict[str, any]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if text.strip():
                return [{
                    'page_number': 1,
                    'text': text.strip(),
                    'file_name': os.path.basename(file_path)
                }]
            else:
                return []
                
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            raise
    
    def process_document(self, file_path: str) -> List[Dict[str, any]]:
        """Process document based on file type"""
        # Validate file first
        is_valid, message = self.validate_file(file_path)
        if not is_valid:
            raise ValueError(message)
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        import re
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep Urdu text
        text = re.sub(r'[^\w\s\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF.,!?;:()\[\]{}"\'-]', '', text)
        
        # Normalize line breaks
        text = text.replace('\n', ' ').replace('\r', ' ')
        
        return text.strip()
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[Dict[str, any]]:
        """Process multiple documents and return combined results"""
        all_pages = []
        
        for file_path in file_paths:
            try:
                pages = self.process_document(file_path)
                # Clean text for each page
                for page in pages:
                    page['text'] = self.clean_text(page['text'])
                all_pages.extend(pages)
                logger.info(f"Successfully processed: {file_path}")
                
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {str(e)}")
                raise
        
        logger.info(f"Total pages processed: {len(all_pages)}")
        return all_pages
